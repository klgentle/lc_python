import datetime
import logging
import os
import re
import sys
import time
import platform

logging.basicConfig(level=logging.DEBUG, format="%(levelname)s: %(message)s")
# 绝对路径的import
sys.path.append(os.path.dirname(os.path.abspath(__file__)) + "/../")

from automatic_office.Holiday import Holiday
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


class CheckInForm(object):
    """
    word style
    """

    def __init__(self, year_month: str):
        # date calculate # 本月5号到下月4号
        if len(year_month) != 6:
            print("Please input year_month with format: YYYYMM!")
            sys.exit(1)
        self.__input_year = int(year_month[:4])
        self.__input_month = int(year_month[4:])
        self.__date_start = self.calculate_date_start()
        self.__date_end = self.calculate_date_end()

        self.__work_date_start = self.__date_start.strftime("%Y-%m-%d")
        self.__work_date_end = self.__date_end.strftime("%Y-%m-%d")
        self.__year_month = year_month

    def calculate_date_start(self) -> datetime.date:
        return datetime.date(self.__input_year, self.__input_month, 5)

    def calculate_date_end(self) -> datetime.date:
        year = self.__input_year
        month = self.__input_month + 1
        if month == 13:
            month = 12
            year += 1
        return datetime.date(year, month, 4)

    def get_all_nature_date(self) -> list:
        """
        返回所有日期列表
        """
        all_nature_date = []
        from_date = self.__date_start
        while from_date.__le__(self.__date_end):
            all_nature_date.append(from_date)
            from_date += datetime.timedelta(1)

        return all_nature_date

    def get_all_work_date(self) -> list:
        """
        返回所有工作日列表
        """
        all_work_date = []
        from_date = self.__date_start
        holiday_obj = Holiday()
        while from_date.__le__(self.__date_end):
            # time.sleep(1)
            # skip holiday
            if not holiday_obj.is_holiday(from_date.strftime("%Y%m%d")):
                all_work_date.append(from_date)
            from_date += datetime.timedelta(1)
            #logging.debug(
            #    "debug work date --{0}".format(from_date.strftime("%Y%m%d"))
            #)

        # TODO write work_date
        return all_work_date

    def check_in_add_table(self, document: object, form_type: str):
        head_list = ["序号", "日期", "签到时间", "签退时间", "备注"]
        table = document.add_table(rows=1, cols=len(head_list), style="Table Grid")
        hdr_cells = table.rows[0].cells
        for i in range(len(head_list)):
            hdr_cells[i].text = head_list[i]
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 水平居中
            hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 垂直居中
            hdr_cells[i].width = Cm(2.82)
        date_list = []
        if form_type.lower() == "normal":
            date_list = self.get_all_work_date()
        elif form_type.lower() == "overtime":
            date_list = self.get_all_nature_date()

        for ind, date in enumerate(date_list):
            row_cells = table.add_row().cells
            row_cells[0].text = str(ind + 1)
            row_cells[1].text = date.strftime("%Y-%m-%d")
            # TODO set font name
            # font.name = u'Calibri'
            # run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Calibri')
            row_cells[2].text = ""
            row_cells[3].text = ""
            row_cells[4].text = ""

            # 设置水平居中
            for i in range(len(head_list)):
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 水平居中
                row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 垂直居中
                row_cells[i].width = Cm(2.82)
        # 设置行高
        for row in table.rows:
            if form_type.lower() == "normal":
                row.height = Cm(0.88)
            elif form_type.lower() == "overtime":
                row.height = Cm(0.70)
        return table

    def write_form(self, form_type: str):
        form_type_name_dict = {"overtime": "带加班", "normal": "正常"}
        form_type_head_dict = {"overtime": "签到表（含加班/调休/请假）", "normal": "签到表"}

        document = Document()
        document.styles["Normal"].font.name = u"宋体"  # 微软雅黑
        document.styles["Normal"].font.size = Pt(10.5)
        document.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
        # 页面布局
        sec = document.sections[0]  # sections对应文档中的“节”
        sec.top_margin = Cm(2.3)
        sec.bottom_margin = Cm(2.3)
        sec.page_height = Cm(29.7)
        sec.page_with = Cm(21)

        # add header
        paragraph = document.add_paragraph()
        run = paragraph.add_run(form_type_head_dict.get(form_type))
        font = run.font
        # set font name
        # font.name = u'宋体'
        # run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        font.size = Pt(16)  # 三号
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_format = paragraph.paragraph_format
        # paragraph_format.line_spacing = 1.5 # 1.5倍行间距
        # paragraph_format.space_after =Pt(0)       #设置段后间距

        paragraph = document.add_paragraph(
            "姓名： 			 日期：{} 至 {}".format(self.__work_date_start, self.__work_date_end)
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Inches(0.3)  # 首行缩进
        paragraph_format.left_indent = Cm(3.7)  # 左侧缩进3.7cm
        # paragraph_format.line_spacing = 1.5 # 1.5倍行间距
        paragraph_format.space_after = Pt(0)  # 设置段后间距

        table = self.check_in_add_table(document, form_type)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        document.save(
            os.path.join(
                "doc_file",
                "文思员工-{0}月签到表_新模版_{1}.docx".format(
                    self.__input_month, form_type_name_dict.get(form_type)
                ),
            )
        )
        logging.info("Done!")


if __name__ == "__main__":
    if len(sys.argv) <= 1:
        logging.info("Please input year_moth(YYYYMM)")
        sys.exit(1)

    obj = CheckInForm(sys.argv[1])
    obj.write_form("overtime")
    obj.write_form("normal")
