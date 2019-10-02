import unittest
import docx


class TestDocx(unittest.TestCase):
    def test_para_run(self):
        doc = docx.Document('automatic_office/文思员工-8月签到表_新模版_带加班.docx')
        # print(len(doc.paragraphs))
        print(f"styles:{doc.styles}")
        for style in doc.styles:
            print(f"doc style name:{style.name}")

        for ind, para in enumerate(doc.paragraphs):
           print(f"{ind}: {para.text}, style name:{para.style.name}")

        for table in doc.tables:
            print(f"table style name0:{table.style.name}")
            # print 第二列的内容，包括表头
            #for cell in table.column_cells(1):
            #    #print(cell.text)
            #    print(f"table style name:{cell.style.name}")


if __name__ == "__main__":
    unittest.main()
